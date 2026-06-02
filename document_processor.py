import os
from typing import List
from PyPDF2 import PdfReader
from docx import Document

class DocumentProcessor:
    """
    文档处理类：负责读取和提取PDF、DOCX文档中的文本
    AI生成：使用Trae辅助编写
    """
    
    def __init__(self):
        self.supported_extensions = ['.pdf', '.docx']
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """
        从PDF文件中提取文本
        AI生成：使用Trae辅助编写
        """
        text = ""
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            print(f"读取PDF文件 {file_path} 时出错: {str(e)}")
        return text
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """
        从DOCX文件中提取文本
        AI生成：使用Trae辅助编写
        """
        text = ""
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
        except Exception as e:
            print(f"读取DOCX文件 {file_path} 时出错: {str(e)}")
        return text
    
    def extract_text(self, file_path: str) -> str:
        """
        根据文件扩展名自动选择提取方法
        AI生成：使用Trae辅助编写
        """
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif ext == '.docx':
            return self.extract_text_from_docx(file_path)
        else:
            print(f"不支持的文件格式: {ext}")
            return ""
    
    def process_directory(self, directory: str) -> List[dict]:
        """
        批量处理指定文件夹内的所有文档
        返回包含文件名和文本内容的字典列表
        AI生成：使用Trae辅助编写
        """
        documents = []
        
        if not os.path.exists(directory):
            print(f"目录不存在: {directory}")
            return documents
        
        for filename in os.listdir(directory):
            ext = os.path.splitext(filename)[1].lower()
            if ext in self.supported_extensions:
                file_path = os.path.join(directory, filename)
                print(f"正在处理: {filename}")
                text = self.extract_text(file_path)
                if text.strip():
                    documents.append({
                        "filename": filename,
                        "content": text,
                        "source": file_path
                    })
                    print(f"  - 提取到 {len(text)} 个字符")
                else:
                    print(f"  - 警告: 文件内容为空")
        
        print(f"\n共处理 {len(documents)} 个文档")
        return documents
    
    def process_uploaded_files(self, uploaded_files) -> List[dict]:
        """
        处理Streamlit上传的文件
        AI生成：使用Trae辅助编写
        """
        documents = []
        
        for uploaded_file in uploaded_files:
            filename = uploaded_file.name
            ext = os.path.splitext(filename)[1].lower()
            
            if ext in self.supported_extensions:
                print(f"正在处理上传文件: {filename}")
                
                if ext == '.pdf':
                    text = ""
                    try:
                        reader = PdfReader(uploaded_file)
                        for page in reader.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                    except Exception as e:
                        print(f"读取PDF时出错: {str(e)}")
                        
                elif ext == '.docx':
                    text = ""
                    try:
                        doc = Document(uploaded_file)
                        for paragraph in doc.paragraphs:
                            if paragraph.text.strip():
                                text += paragraph.text + "\n"
                    except Exception as e:
                        print(f"读取DOCX时出错: {str(e)}")
                else:
                    text = ""
                
                if text.strip():
                    documents.append({
                        "filename": filename,
                        "content": text,
                        "source": filename
                    })
        
        return documents


if __name__ == "__main__":
    processor = DocumentProcessor()
    test_dir = "./documents"
    if os.path.exists(test_dir):
        docs = processor.process_directory(test_dir)
        for doc in docs:
            print(f"\n文件: {doc['filename']}")
            print(f"内容预览: {doc['content'][:200]}...")